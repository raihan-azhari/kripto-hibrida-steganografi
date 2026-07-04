from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION_START

ROOT = Path(__file__).resolve().parents[1]
OUT = Path('/mnt/data/Laporan_Kripto_Hibrida_Raihan_Azhari_Lubis.docx')


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = 'Times New Roman'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    r.font.size = Pt(10)
    if color:
        r.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        r.font.color.rgb = RGBColor(30, 60, 100)
    return p


def add_para(doc, text='', align=None, bold_start=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    if align:
        p.alignment = align
    if bold_start and text.startswith(bold_start):
        r = p.add_run(bold_start)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
        r2 = p.add_run(text[len(bold_start):])
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    return p


def add_code(doc, code, max_lines=None):
    lines = code.splitlines()
    if max_lines:
        lines = lines[:max_lines]
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(line if line else ' ')
        r.font.name = 'Courier New'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
        r.font.size = Pt(7.5)
    return


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], '1F4E79')
        set_cell_text(hdr[i], h, bold=True, color=(255, 255, 255))
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val))
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_image(doc, path, width=6.2, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(9)
    return p


def build_report():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)

    styles = doc.styles
    styles['Normal'].font.name = 'Times New Roman'
    styles['Normal'].font.size = Pt(11)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('LAPORAN TUGAS PROGRAM\nKRIPTOGRAFI HIBRIDA + STEGANOGRAFI')
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(30,60,100)
    doc.add_paragraph('\n')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Implementasi AES, RSA, DSA, SHA-256, Diffie-Hellman, dan Modified LSB')
    r.font.name = 'Times New Roman'; r.font.size = Pt(14)
    doc.add_paragraph('\n\n')
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    cover_rows = [('Nama', 'Raihan Azhari Lubis'), ('NIM', '231111619'), ('Jenis Tugas', 'Program Kripto Hibrida + Steganografi')]
    for i, (k,v) in enumerate(cover_rows):
        set_cell_text(table.rows[i].cells[0], k, bold=True)
        set_cell_text(table.rows[i].cells[1], v)
    doc.add_paragraph('\n')
    add_para(doc, 'Dokumen ini berisi uraian pekerjaan, desain sistem, cara menjalankan program, bukti eksekusi, dan source code inti untuk memenuhi instruksi pengumpulan tugas.', align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    add_heading(doc, 'Ringkasan', 1)
    add_para(doc, 'Program yang dibuat adalah sistem pengamanan file berbasis kriptografi hibrida dan steganografi. File plaintext dienkripsi memakai AES-256-GCM. Kunci AES dilindungi dengan RSA-OAEP-SHA256. Paket terenkripsi ditandatangani dengan DSA-SHA256. SHA-256 digunakan untuk hash file asli dan hash material tanda tangan. Diffie-Hellman dengan HKDF-SHA256 digunakan untuk menurunkan seed posisi penyisipan steganografi. Payload akhir disisipkan ke dalam gambar PNG menggunakan modified randomized LSB dua bit pada byte RGB.')
    add_para(doc, 'Hasil demo otomatis menunjukkan pertukaran kunci Diffie-Hellman cocok, tanda tangan DSA valid, hash file asli sama dengan hash file hasil dekripsi, dan konten file pulih sama persis. Dengan demikian, program memenuhi unsur kerahasiaan, integritas, autentikasi, pertukaran kunci, dan penyembunyian data pada citra.')

    add_heading(doc, 'Kesesuaian dengan Instruksi Tugas', 1)
    add_table(doc, ['Komponen Instruksi', 'Implementasi pada Program', 'Status'], [
        ('AES', 'AES-256-GCM untuk mengenkripsi isi file.', 'Terpenuhi'),
        ('Algoritma kunci asimetris untuk pengamanan kunci', 'RSA-OAEP-SHA256 untuk mengenkripsi kunci AES.', 'Terpenuhi'),
        ('Tanda tangan digital', 'DSA-SHA256 untuk menandatangani paket terenkripsi.', 'Terpenuhi'),
        ('SHA-256', 'Hash plaintext dan hash material tanda tangan.', 'Terpenuhi'),
        ('Pertukaran Kunci Diffie-Hellman', 'DH + HKDF-SHA256 untuk menurunkan seed posisi LSB.', 'Terpenuhi'),
        ('Steganografi modifikasi LSB', 'Randomized LSB dua bit per byte RGB dengan frame magic dan panjang payload.', 'Terpenuhi'),
        ('Bukti eksekusi', 'Demo otomatis dan log hasil uji disertakan pada laporan.', 'Terpenuhi'),
        ('Cara menjalankan program', 'Langkah instalasi dan perintah CLI disertakan.', 'Terpenuhi'),
        ('Source code', 'Source code inti dicantumkan dalam lampiran dan source lengkap ada di ZIP project.', 'Terpenuhi'),
    ], widths=[1.8, 4.0, 1.0])

    add_heading(doc, 'Analisis Kebutuhan', 1)
    add_para(doc, 'Kebutuhan fungsional sistem adalah menerima file plaintext dan cover image, menghasilkan stego image yang memuat paket terenkripsi, serta mengembalikan file plaintext melalui proses ekstraksi dan dekripsi. Sistem juga harus dapat membuat kunci kriptografi, membuktikan pertukaran kunci Diffie-Hellman, memverifikasi tanda tangan digital, dan mengecek integritas output dengan SHA-256.')
    add_para(doc, 'Kebutuhan nonfungsionalnya adalah program mudah dijalankan dari command line, source code modular, payload terenkripsi tidak mudah dibaca tanpa kunci, perubahan payload dapat terdeteksi, dan perubahan visual pada gambar cover tidak mencolok. Oleh karena itu steganografi tidak dilakukan secara sekuensial biasa, tetapi memakai urutan posisi acak dari seed hasil DH-HKDF.')

    add_heading(doc, 'Desain Arsitektur Sistem', 1)
    add_image(doc, ROOT/'docs/architecture_flow.png', width=6.8, caption='Gambar 1. Alur sistem kriptografi hibrida dan steganografi.')
    add_para(doc, 'Alur utama dimulai dari plaintext. Sistem menghitung hash SHA-256 plaintext, membuat kunci AES acak, lalu mengenkripsi plaintext memakai AES-256-GCM. Kunci AES dienkripsi dengan RSA-OAEP milik penerima. Metadata, ciphertext, encrypted key, nonce, hash, dan algoritma dibentuk menjadi paket JSON. Paket tersebut ditandatangani dengan DSA. Setelah itu, Diffie-Hellman menurunkan seed posisi acak untuk penyisipan LSB. Paket JSON bertanda tangan disisipkan pada cover image sehingga menjadi stego image.')
    add_para(doc, 'Pada proses dekripsi, penerima menurunkan seed yang sama memakai private key DH penerima dan public key DH pengirim. Seed digunakan untuk membaca posisi LSB yang sama. Payload JSON diekstrak, tanda tangan DSA diverifikasi, kunci AES dibuka memakai private key RSA penerima, ciphertext didekripsi dengan AES-GCM, lalu SHA-256 file hasil dekripsi dibandingkan dengan hash asli.')

    add_heading(doc, 'Rincian Algoritma yang Digunakan', 1)
    add_heading(doc, 'AES-256-GCM', 2)
    add_para(doc, 'AES digunakan sebagai algoritma simetris untuk mengenkripsi isi file karena cocok untuk data berukuran besar. Mode GCM dipakai karena menyediakan enkripsi dan autentikasi ciphertext. Program memakai kunci 256 bit dan nonce 96 bit yang dibuat acak setiap proses enkripsi.')
    add_heading(doc, 'RSA-OAEP-SHA256', 2)
    add_para(doc, 'RSA digunakan untuk melindungi kunci AES, bukan untuk mengenkripsi seluruh file. Strategi ini adalah pola hibrida: data besar dienkripsi dengan AES, sementara kunci AES yang kecil diamankan dengan RSA. OAEP dengan SHA-256 dipakai sebagai padding modern untuk mengurangi risiko serangan pada RSA mentah.')
    add_heading(doc, 'DSA-SHA256', 2)
    add_para(doc, 'DSA digunakan untuk tanda tangan digital. Paket JSON dibuat secara canonical, kemudian ditandatangani memakai private key DSA pengirim. Penerima memverifikasi signature memakai public key DSA pengirim. Jika payload berubah satu bit saja, verifikasi tanda tangan akan gagal.')
    add_heading(doc, 'SHA-256', 2)
    add_para(doc, 'SHA-256 digunakan pada dua titik: pertama, menghitung hash plaintext asli; kedua, menghitung hash material tanda tangan. Hash plaintext dibandingkan kembali setelah dekripsi untuk membuktikan bahwa file pulih identik dengan file asli.')
    add_heading(doc, 'Diffie-Hellman + HKDF-SHA256', 2)
    add_para(doc, 'Diffie-Hellman digunakan agar pengirim dan penerima menghasilkan secret yang sama tanpa mengirim private key. Secret mentah dari DH tidak langsung dipakai, tetapi diproses dengan HKDF-SHA256 menjadi seed 32 byte. Seed ini mengendalikan urutan pseudo-random posisi byte carrier untuk modified LSB.')
    add_heading(doc, 'Modified Randomized LSB', 2)
    add_image(doc, ROOT/'docs/lsb_diagram.png', width=6.8, caption='Gambar 2. Ilustrasi penyisipan modified LSB dua bit.')
    add_para(doc, 'Metode LSB yang dipakai dimodifikasi dengan dua cara. Pertama, penyisipan menggunakan dua bit terakhir pada setiap byte RGB sehingga kapasitas lebih besar dibanding satu bit. Kedua, posisi byte carrier diacak memakai seed hasil DH-HKDF, sehingga payload tidak tersimpan pada urutan piksel linear. Paket juga diberi magic bytes KHB1 dan panjang payload 32 bit agar proses ekstraksi dapat memvalidasi format.')

    add_heading(doc, 'Format Paket Terenkripsi', 1)
    add_para(doc, 'Paket yang disisipkan ke gambar berbentuk JSON. Format ini dipilih agar metadata algoritma, nonce, ciphertext, encrypted AES key, hash, dan tanda tangan mudah diaudit. Data biner disimpan dalam Base64.')
    add_table(doc, ['Field', 'Fungsi'], [
        ('version', 'Versi format paket.'),
        ('student', 'Identitas nama dan NIM.'),
        ('algorithm', 'Daftar algoritma yang dipakai.'),
        ('nonce_b64', 'Nonce AES-GCM.'),
        ('encrypted_aes_key_b64', 'Kunci AES yang sudah dienkripsi RSA.'),
        ('ciphertext_b64', 'Hasil enkripsi plaintext memakai AES-GCM.'),
        ('original_sha256', 'Hash file plaintext asli.'),
        ('signature_b64', 'Tanda tangan DSA atas paket.'),
        ('signing_material_sha256', 'Hash canonical JSON yang ditandatangani.'),
    ], widths=[2.0, 4.8])

    add_heading(doc, 'Struktur Folder Project', 1)
    add_code(doc, '''kripto_hibrida_raihan/
|-- main.py
|-- requirements.txt
|-- README.md
|-- GITHUB_UPLOAD_GUIDE.md
|-- samples/
|   |-- cover.png
|   `-- pesan_raihan.txt
|-- docs/
|   |-- demo_execution_log.txt
|   |-- stego_raihan_demo.png
|   `-- recovered_pesan_raihan_demo.txt
`-- src/
    `-- hybrid_crypto/
        |-- crypto_engine.py
        |-- demo_data.py
        |-- key_exchange.py
        |-- key_manager.py
        |-- stego_lsb.py
        |-- utils.py
        `-- __init__.py''')

    add_heading(doc, 'Cara Menjalankan Program', 1)
    add_para(doc, 'Langkah pertama adalah memasang dependency. Program membutuhkan Python 3.10 atau lebih baru, package cryptography, dan Pillow.')
    add_code(doc, '''python -m venv .venv
.venv\\Scripts\\activate
pip install -r requirements.txt''')
    add_para(doc, 'Untuk menjalankan semua proses secara otomatis, gunakan perintah berikut pada folder project:')
    add_code(doc, '''python main.py demo --out demo_run''')
    add_para(doc, 'Perintah demo akan membuat kunci, membuat plaintext dan cover image contoh, melakukan enkripsi, menyisipkan payload ke gambar, mengekstrak payload, mendekripsi file, dan membandingkan hash SHA-256.')
    add_para(doc, 'Perintah manual untuk generate keys:')
    add_code(doc, '''python main.py gen-keys --out keys''')
    add_para(doc, 'Perintah manual untuk enkripsi dan penyisipan ke cover image:')
    add_code(doc, '''python main.py encrypt ^
  --input samples/pesan_raihan.txt ^
  --cover samples/cover.png ^
  --output output/stego_raihan.png ^
  --recipient-rsa-public keys/receiver_rsa_public.pem ^
  --sender-dsa-private keys/sender_dsa_private.pem ^
  --sender-dh-private keys/sender_dh_private.pem ^
  --recipient-dh-public keys/receiver_dh_public.pem''')
    add_para(doc, 'Perintah manual untuk ekstraksi dan dekripsi:')
    add_code(doc, '''python main.py decrypt ^
  --stego output/stego_raihan.png ^
  --output output/recovered_pesan_raihan.txt ^
  --recipient-rsa-private keys/receiver_rsa_private.pem ^
  --sender-dsa-public keys/sender_dsa_public.pem ^
  --sender-dh-public keys/sender_dh_public.pem ^
  --recipient-dh-private keys/receiver_dh_private.pem''')

    add_heading(doc, 'Bukti Eksekusi', 1)
    add_para(doc, 'Demo otomatis sudah dijalankan pada lingkungan kerja. Bukti eksekusi menunjukkan match Diffie-Hellman bernilai true, signature_status bernilai valid, dan same_content bernilai true. Artinya secret DH sama, tanda tangan digital valid, dan file hasil dekripsi identik dengan file awal.')
    add_image(doc, ROOT/'docs/console_log_part_1.png', width=6.8, caption='Gambar 3. Bukti eksekusi bagian 1: pembuatan kunci dan file contoh.')
    add_image(doc, ROOT/'docs/console_log_part_2.png', width=6.8, caption='Gambar 4. Bukti eksekusi bagian 2: DH, enkripsi, dan steganografi.')
    add_image(doc, ROOT/'docs/console_log_part_3.png', width=6.8, caption='Gambar 5. Bukti eksekusi bagian 3: dekripsi dan verifikasi akhir.')
    add_para(doc, 'Nilai hash SHA-256 plaintext dan recovered file pada demo sama, yaitu e48688c8f189652f11ba320242937f6b9424188e3638e508240982f83c3d8466. Nilai same_content true membuktikan bahwa isi file hasil dekripsi tidak berubah.')
    add_image(doc, ROOT/'docs/cover_vs_stego.png', width=6.8, caption='Gambar 6. Cover image dan stego image terlihat sama secara visual.')

    add_heading(doc, 'Analisis Keamanan dan Keterbatasan', 1)
    add_para(doc, 'Kerahasiaan isi file dijaga oleh AES-256-GCM. Tanpa kunci AES yang benar, ciphertext tidak dapat didekripsi. Kunci AES tidak disimpan mentah, tetapi dienkripsi dengan RSA-OAEP public key penerima. Dengan demikian hanya pemilik private key RSA penerima yang dapat membuka kunci AES.')
    add_para(doc, 'Integritas dan autentikasi dijaga oleh dua lapis mekanisme. AES-GCM memiliki tag autentikasi terhadap ciphertext, sedangkan DSA-SHA256 memverifikasi bahwa paket tidak diubah dan berasal dari pemilik private key DSA. SHA-256 plaintext juga menjadi pembanding akhir setelah proses dekripsi.')
    add_para(doc, 'Diffie-Hellman dipakai sebagai faktor tambahan untuk steganografi. Jika private/public key DH yang dipakai tidak sesuai, urutan posisi LSB berbeda sehingga header KHB1 tidak terbaca. Artinya payload tersembunyi tidak mudah ditemukan hanya dengan membaca LSB secara berurutan.')
    add_para(doc, 'Keterbatasan program adalah hanya mendukung output stego PNG agar perubahan LSB tidak rusak oleh kompresi lossy. Jika stego image dikonversi ke JPEG atau diubah ukurannya, payload kemungkinan rusak. Program ini juga ditujukan untuk pembelajaran akademik, sehingga manajemen kunci produksi seperti passphrase private key, rotasi kunci, dan audit akses belum dibuat lengkap.')

    add_heading(doc, 'Kesimpulan', 1)
    add_para(doc, 'Program berhasil mengimplementasikan kombinasi kriptografi hibrida dan steganografi sesuai instruksi. Seluruh komponen yang diminta tersedia: AES, RSA, DSA, SHA-256, Diffie-Hellman, dan modified LSB. Bukti eksekusi menunjukkan bahwa proses enkripsi, penyisipan, ekstraksi, dekripsi, verifikasi tanda tangan, dan validasi hash berjalan benar pada data uji.')
    add_para(doc, 'Source code sudah dibuat modular agar mudah diperiksa dan diunggah ke GitHub. File README.md dan GITHUB_UPLOAD_GUIDE.md disediakan untuk membantu proses menjalankan dan mengunggah project.')

    add_heading(doc, 'Lampiran A - Source Code Inti', 1)
    add_para(doc, 'Bagian ini memuat source code inti. Source code lengkap, sample file, README, dan panduan GitHub juga disertakan pada ZIP project.')

    code_files = [
        'main.py',
        'src/hybrid_crypto/crypto_engine.py',
        'src/hybrid_crypto/stego_lsb.py',
        'src/hybrid_crypto/key_manager.py',
        'src/hybrid_crypto/key_exchange.py',
        'src/hybrid_crypto/utils.py',
        'src/hybrid_crypto/demo_data.py',
    ]
    for rel in code_files:
        doc.add_page_break()
        add_heading(doc, f'Listing: {rel}', 2)
        code = (ROOT/rel).read_text(encoding='utf-8')
        add_code(doc, code)

    # Footer
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = footer.add_run('Raihan Azhari Lubis - 231111619 - Kriptografi Hibrida + Steganografi')
        r.font.name = 'Times New Roman'
        r.font.size = Pt(9)

    doc.save(OUT)
    print(OUT)

if __name__ == '__main__':
    build_report()
